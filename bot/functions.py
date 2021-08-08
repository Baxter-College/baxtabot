# functions.py
#
# Includes all non-messaging functionality of baxterbot
import traceback

import datetime
from dateutil.parser import parse
import random
import json
import requests
import math
import mammoth
import re
import yagmail
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from bot.settings import *

import bot.models as models
import bot.extract as extract


# ====== Specific functions ===== #
def findMeal(message):
    if "dinner" in message:
        meal = "dinner"
    elif "lunch" in message:
        meal = "lunch"
    elif "breakfast" in message:
        meal = "breakfast"
    else:
        return False
    return meal


def findTime(message):

    addTime = datetime.timedelta(hours=0)

    today = datetime.datetime.today() + datetime.timedelta(hours=11)
    days = {
        "monday": 0,
        "tuesday": 1,
        "wednesday": 2,
        "thursday": 3,
        "friday": 4,
        "saturday": 5,
        "sunday": 6,
    }

    if "tomorrow" in message or "tommorow" in message:
        addTime += datetime.timedelta(hours=24)

    for day in days:
        if day in message:  # if the day is mentioned
            dayDiff = days[day] - today.weekday()  # get difference between days
            if dayDiff < 0:  # if the day is "behind" current day, make day next week
                dayDiff += 7
            elif "next" in message:
                addTime += datetime.timedelta(hours=24 * 7)  # add this day next week

            addTime += datetime.timedelta(hours=24 * dayDiff)

    return addTime


def dinoRequest(meal, addTime):
    # meal is "dinner", "lunch" or "breakfast"

    ten_hours = datetime.timedelta(hours=11)

    today = datetime.datetime.now()

    today_AEST = today + ten_hours

    today_AEST += addTime  # if no add time, timedelta will be 0 hours so no effect

    print("Date is: {}".format(today_AEST.date().strftime("%Y-%m-%d")))

    try:
        dino = (
            models.Meal.select()
            .where(models.Meal.date == today_AEST.date())
            .where(models.Meal.type == meal)
            .get()
        )
    except Exception as e:
        print("---> ", e)
        print("Meal: ", meal)
        print("Date: ", today_AEST.date())
        return "Honestly.... I don't know."

    return "{} at dino is:\n{}".format(meal, dino.description)

def getTimeFromAddTime(addTime):
    ten_hours = datetime.timedelta(hours=11)

    today = datetime.datetime.now()

    today_AEST = today + ten_hours

    today_AEST += addTime  # if no add time, timedelta will be 0 hours so no effect

    return today_AEST

def dinoRequestObj(meal, addTime):
    # meal is "dinner", "lunch" or "breakfast"
    today_AEST = getTimeFromAddTime(addTime)
    print("Date is: {}".format(today_AEST.date().strftime("%Y-%m-%d")))

    try:
        dino = (
            models.Meal.select()
            .where(models.Meal.date == today_AEST.date())
            .where(models.Meal.type == meal)
            .get()
        )
    except Exception as e:
        return None

    return dino


def makeDinoVote(vote):

    dino = getCurrentDino()

    if vote == "goodvote":
        print("the meal has: {} likes".format(dino.likes))
        dino.likes = dino.likes + 1
    elif vote == "badvote":
        print("the meal has: {} likes".format(dino.dislikes))
        dino.dislikes = dino.likes + 1

    dino.save()


def dinoPoll():

    dino = getCurrentDino()

    if dino.likes == 0 and dino.dislikes == 0:
        message = "No one has voted for {}! 😢\nYou can be the first to vote with 'dinovote'".format(
            dino.type
        )

    elif dino.likes < dino.dislikes:
        perc = (dino.dislikes / (dino.dislikes + dino.likes)) * 100
        message = "{}% of people disliked {}.".format(perc, dino.type)

    elif dino.likes > dino.dislikes:
        perc = (dino.likes / (dino.dislikes + dino.likes)) * 100
        message = "{}% of people enjoyed {}!!!".format(perc, dino.type)

    else:
        message = "The crowd is split! Dino is a polarising meal.\nLet me know your thoughts with 'dinovote'"

    return message


def getCurrentDino(time=None):

    if time is None:
        time = datetime.datetime.now() + datetime.timedelta(hours=11)  # to make it aest

    today = datetime.datetime.today() + datetime.timedelta(hours=11)
    breakfast = today.replace(hour=7, minute=0)
    lunch = today.replace(hour=12, minute=0)
    dinner = today.replace(
        hour=16, minute=0
    )  # just to make sure (starting at 4 so people can ask what's dino earlier for dinner)

    try:
        if time < lunch:
            # for today's breakfast
            dino = (
                models.Meal.select()
                .where(models.Meal.type == "breakfast")
                .where(models.Meal.date == today.date())
                .get()
            )
        elif time >= lunch and time < dinner:
            # for today's lunch
            dino = (
                models.Meal.select()
                .where(models.Meal.type == "lunch")
                .where(models.Meal.date == today.date())
                .get()
            )
        elif time >= dinner:
            # for today's dinner
            dino = (
                models.Meal.select()
                .where(models.Meal.type == "dinner")
                .where(models.Meal.date == today.date())
                .get()
            )
        return dino
    except:
        return None

# ======== Late Meals ======= #

def orderLateMeal(message, sender_psid):
    meal_name = findMeal(message)
    if not meal_name:
        addTime = datetime.timedelta(hours=0)
        meal = getCurrentDino()
        meal_name = meal.type
    else:
        addTime = findTime(message)
        meal = dinoRequestObj(meal_name, addTime)

    if meal is None:
        raise Exception('Meal does not exist - dino menu needs updating')
    meal_id = meal.id

    ressie = getRessieBySender(sender_psid).id
    notes = 'See dietary requirement records'

    models.LateMeal.create(meal=meal_id, ressie=ressie, notes=notes, completed=False)

    return meal_name, getTimeFromAddTime(addTime).date().strftime('%d/%m/%Y')


def generateStickersDocument(oustanding_meals):
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(0.46)
        section.right_margin = Cm(0.46)

    meals_processed = 0
    print(oustanding_meals)

    while True:
        if meals_processed >= len(oustanding_meals):
            break

        table = document.add_table(rows=7, cols=2)

        for row in table.rows:
            row.height = Cm(3.76)
            for cell in row.cells:

                if meals_processed >= len(oustanding_meals):
                    break

                meal = oustanding_meals[meals_processed]
                print(meal)
                cell.text = f"\n{meal['id']} {meal['first_name']} {meal['last_name']}\n{meal['college']}\n{meal['date']}\n{meal['dietaries']}"
                cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                meals_processed += 1

    print('Meal processing completed')
    document.save('LatemealStickers/LatemealStickers.docx')


def sendLateMealStickersEmail():
    email = yagmail.SMTP('baxtabot21@gmail.com', 'meqdeh-5Jysve-xewtuc')
    email.send('n.patrikeos@student.unsw.edu.au', 'Late meals', contents='LatemealStickers/LatemealStickers.docx')

def generateLateMealStickers(meals):

    oustanding_meals = models.LateMeal.select(models.LateMeal.id, models.Ressie.first_name, models.Ressie.last_name, models.Ressie.college, models.Meal.date, models.Meal.type, models.Client.dietaries).join(models.Ressie).join(models.Client).switch(models.LateMeal).join(models.Meal).where(models.LateMeal.id << meals).dicts()
    generateStickersDocument(oustanding_meals)
    sendLateMealStickersEmail()

def getRessieBySender(sender_psid):
    data = humanisePSID(sender_psid)
    print(data)
    if not data:
        print("received message from ghost!")
        return

    name = data['first_name'] + ' ' + data['last_name']

    _, confidence, ressie = models.Ressie.fuzzySearch(name)
    if confidence <= 70:
        print(sender_psid)
        if sender_psid != 'cmd':
            raise Exception('Ressie not found')
    else:
        return ressie


# ======== J&D ========== #



# ===== Shopen ===== #

# TODO: integrate this toggle action into a function so we are not duplicating functionality


def set_shop(rs, switch):

    if switch[0].lower() == "on":
        message.bot.set_variable("shop", True)
        return "Shopen!"
    else:
        message.bot.set_variable("shop", None)
        return "Shclosed 😭"


def get_shop(rs, args):

    shop = message.bot.get_variable("shop")

    return "Shopen!!!" if shop else "Shclosed 😭"


# ===== Baxter Events ===== #
# TODO: Move this into message module
def uploadAsset(assetUrl):

    r = requests.post(
        "https://graph.facebook.com/v2.6/me/message_attachments",
        params={"access_token": PAGE_ACCESS_TOKEN},
        json={
            "message": {
                "attachment": {
                    "type": "image",
                    "payload": {"is_reusable": True, "url": assetUrl},
                }
            }
        },
    )

    if r.status_code == 200:
        return r.json()
    else:
        print("Asset Upload has gone to shit -> ", r.status_code)
        return None


# ===== Hashbrowns ===== #

def set_hashbrowns(rs, switch):

    message.bot.set_variable('hashbrownsday', datetime.date.today())
    if switch[0].lower() == 'on':
        message.bot.set_variable('hashbrowns', True)
        return "OMG best news ever! 😃 Your friends will arrive shortly..."
    else:
        message.bot.set_variable('hashbrowns', None)
        return "N-n-n-noooooooo! 😭 Enjoy a lonely Dino, knowing you took one for the team..."

def get_hashbrowns(rs, args):

    if (message.bot.get_variable('hashbrownsday') == datetime.date.today()):
        hashbrowns = message.bot.get_variable('hashbrowns')
        return "Get out of bed! HASHBROWNS TODAY!!! 🥔🥔🎊🎉🎊🎉" if hashbrowns else "Bad news: no hashbrowns... stay in bed 😔"
    else:
        return "Nobody's been game to find out yet 🤔 Type 'sethashbrowns on' or 'sethashbrowns off' if you happen to get out of bed"


# ===== Week Events ===== #
def getWeekEvents():

    today = datetime.datetime.today() + datetime.timedelta(hours=11)  # to make it aest
    # Take todays date. Subtract the number of days which already passed this week (this gets you 'last' monday).
    week_monday = today + datetime.timedelta(days=-today.weekday(), weeks=0)

    try:
        weekCal = (
            models.WeekCal.select()
            .where(models.WeekCal.week_start == week_monday.date())
            .get()
        )

        return weekCal.assetID
    except:
        return None


# ===== Get Room Number ===== #
def extractName(msg):
    half = msg.split("is", 1)[1].split()
    return " ".join(half[:2])


def getRoomNumber(name):

    try:
        gotName, confidence, ressie = models.Ressie.fuzzySearch(name)
        client = models.Client.select().join(models.Ressie).where(models.Ressie.id == ressie.id)
        if client:
            client = client.get()

            if not client.roomshown:
                return '{} is in baxter and has turned room sharing off'.format(gotName)

        if confidence < 85:
            return "{} is in room {} (I'm {} percent sure I know who you're talking about)".format(
                gotName, ressie.room_number, confidence
            )
        return "{} is in room {}".format(gotName, ressie.room_number)
    except Exception as e:
        traceback.print_exc()
        return """I could not find a room number for '{}' ... are you sure they go to Baxter?\n
        They may not have registered an account at baxtabot.herokuapp.com !! So go nag them to do that xD\n
          \nPlease make sure you spell their full name correctly.\n\n (Fun fact: Some people use names that are not in fact their names. Nicknames won't work)""".format(
            "".join(name).title()
        )

def dinoparse(lines):
    lines = extract.text_replace(lines)

    soup = BeautifulSoup(lines, features="html.parser")
    assert soup != None
    pretty = soup.prettify()

    rows = extract.get_rows(soup)

    mealsByDay = extract.get_meals(rows[1:])
    date, sucess = extract.extract_date(soup)
    return [date, sucess, mealsByDay, pretty]

def extractRessieFromCSV(row):
    first_name, last_name, room_number = row
    last_name = last_name.capitalize()
    last_name_edited = ''
    for char in last_name:
        if char == '(' or char == ' ':
            break
        else:
            last_name_edited += char

    return first_name, last_name, room_number[:3]


def humanisePSID(PSID):
    url = "https://graph.facebook.com/" + str(PSID)
    print("url\n", url)
    print(PAGE_ACCESS_TOKEN)
    r = requests.get(
        url,
        params={
            "fields": "first_name,last_name,profile_pic",
            "access_token": PAGE_ACCESS_TOKEN,
        },
    )

    if r.status_code == 200:
        data = r.json()
        print("Worked!")
        return data
    else:
        print("response")
        print(r.content)
        print(r.status_code)
        print("FUCKED! PSID was: {}".format(str(PSID)))

def createRessie(first_name, last_name, room_number):
    models.Ressie.create(
        first_name = first_name,
        last_name = last_name,
        room_number = room_number,
        floor = int(str(room_number)[:1])
        )  # get the first digit of the room number and set that as floor

def validateTokenPermissions(token, page):
    userperms = models.Client.select(models.ClientPermissions.dinoread, models.ClientPermissions.dinowrite, models.ClientPermissions.ressies,
                                    models.ClientPermissions.calendar, models.ClientPermissions.sport, models.ClientPermissions.latemeals, models.ClientPermissions.users).join(models.ActiveTokens).switch(models.Client).join(models.ClientPermissions).where(models.ActiveTokens.token == token).dicts()

    if userperms is None:
        return False

    userperms = userperms.get()

    if page == 'dinoread' and userperms['dinoread']:
        return True
    elif page == 'dinowrite' and userperms['dinowrite']:
        return True
    elif page == 'ressies' and userperms['ressies']:
        return True
    elif page == 'calendar' and userperms['calendar']:
        return True
    elif page == 'sport' and userperms['sport']:
        return True
    elif page == 'latemeals' and userperms['latemeals']:
        return True
    elif page == 'users' and userperms['users']:
        return True

    return False

def deleteActiveToken(token):
    instance = models.ActiveTokens.select().where(models.ActiveTokens.token == token).get()
    instance.delete_instance()
